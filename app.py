from dotenv import load_dotenv
import os

load_dotenv()  
import re
import json
import tempfile
import os
from PyPDF2 import PdfReader
from docx import Document
import google.generativeai as genai
import streamlit as st

# ==============================
# 📌 Step 1: Setup Gemini AI
# ==============================
genai.configure(api_key=os.getenv("GENAI_API_KEY"))
model = genai.GenerativeModel("gemini-2.0-flash")

# ==============================
# 📌 Step 2: Resume Readers
# ==============================
def read_pdf(file_path):
    try:
        from pdfminer.high_level import extract_text
        return extract_text(file_path)
    except:
        reader = PdfReader(file_path)
        return "\n".join([page.extract_text() or "" for page in reader.pages])

def extract_links_from_pdf(file_path):
    links = []
    reader = PdfReader(file_path)
    for page in reader.pages:
        if "/Annots" in page:
            for annot in page["/Annots"]:
                obj = annot.get_object()
                if "/A" in obj and "/URI" in obj["/A"]:
                    links.append(obj["/A"]["/URI"])
    return list(set(links))

def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs])

def extract_links_from_docx(file_path):
    links = []
    doc = Document(file_path)
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype:
            links.append(rel.target_ref)
    return list(set(links))

def load_resume(resume_file):
    suffix = "." + resume_file.name.split(".")[-1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(resume_file.read())
        tmp_path = tmp.name

    if tmp_path.endswith(".pdf"):
        resume_text = read_pdf(tmp_path)
        links = extract_links_from_pdf(tmp_path)
    elif tmp_path.endswith(".docx"):
        resume_text = read_docx(tmp_path)
        links = extract_links_from_docx(tmp_path)
    else:
        raise ValueError("Unsupported file type. Use PDF or DOCX.")

    return resume_text, links

# ==============================
# 📌 Step 3: Parse Resume with Gemini AI
# ==============================
def parse_resume_with_gemini(resume_text, links_from_file):
    prompt = f"""
You are an intelligent resume parser.

Task:
Extract structured information from the provided resume text and hyperlinks.
Return strictly valid JSON with keys:
- name, email, phone, location
- links: {{linkedin, github, portfolio, others}}
- summary
- skills
- education: list of {{degree, institution, year, score}}
- experience: list of {{company, role, duration, description}}
- projects: list of {{title, description, technologies}}
- certifications
- achievements

Resume Text:
{resume_text}

Extracted Hyperlinks:
{links_from_file}

Instructions:
1. Ensure JSON is valid.
2. Deduplicate skills and links.
3. Return lists where applicable.
"""
    try:
        response = model.generate_content(prompt)
        extracted_data = response.text.strip()
    except Exception as e:
        st.error(f"Gemini API error: {e}")
        return {}

    # Clean Markdown fences
    extracted_data = re.sub(r"^```json", "", extracted_data)
    extracted_data = re.sub(r"^```", "", extracted_data)
    extracted_data = re.sub(r"\s*```$", "", extracted_data)

    try:
        data = json.loads(extracted_data)
    except:
        data = {}

    if "skills" in data and isinstance(data["skills"], list):
        data["skills"] = sorted(list(set(data["skills"])))
    return data

# ==============================
# 📌 Step 4: Load Categorized Roles
# ==============================
def load_categories_roles_skills_tools_from_json(file_path="Technical_Roles.json"):
    with open(file_path, "r", encoding="utf-8") as f:
        categorized_data = json.load(f)
    return categorized_data

# ==============================
# 📌 Step 5: Candidate Type Detection
# ==============================
def detect_candidate_type(resume_data, resume_text):
    experience = resume_data.get("experience", [])
    keywords = ["intern", "internship", "trainee", "fresher"]
    if not experience or (any(k in resume_text.lower() for k in keywords) and len(experience) <= 1):
        return "fresher"
    else:
        return "experienced"

# ==============================
# 📌 Step 6: ATS Scoring
# ==============================
def evaluate_resume_sections_and_score(resume_data, role_info, candidate_type):
    report = {}
    score = 0
    # --- Summary ---
    summary = resume_data.get("summary", "")
    if not summary: report["Summary"] = "❌ Missing"
    elif len(summary) > 30: report["Summary"], score = "✅ Good", score+10
    else: report["Summary"], score = "⚠️ Improve", score+5

    # --- Skills ---
    resume_skills = resume_data.get("skills", [])
    role_skills = role_info.get("skills", [])
    matched_skills = [s for s in resume_skills if s.lower() in map(str.lower, role_skills)]
    if role_skills:
        if len(matched_skills) == 0: report["Skills"] = "❌ Present but not relevant"
        else:
            percent = len(matched_skills)/len(role_skills)*100
            if percent >= 80: report["Skills"], score = "✅ Good", score+30
            else: report["Skills"], score = "⚠️ Improve", score+(percent/100*30)
    else: report["Skills"] = "❌ Missing"

    # --- Tools ---
    role_tools = role_info.get("tools", [])
    matched_tools = [s for s in resume_skills if s.lower() in map(str.lower, role_tools)]
    if role_tools:
        if len(matched_tools) == 0: report["Tools"] = "❌ Present but not relevant"
        else:
            percent_tools = len(matched_tools)/len(role_tools)*100
            if percent_tools >= 80: report["Tools"], score = "✅ Good", score+10
            else: report["Tools"], score = "⚠️ Improve", score+(percent_tools/100*10)
    else: report["Tools"] = "❌ Missing"

    # --- Projects ---
    projects = resume_data.get("projects", [])
    min_projects = 3 if candidate_type=="fresher" else 5
    role_keywords = [x.lower() for x in role_skills + role_tools]
    def project_related(proj): return any(kw in (proj.get("title","")+" "+proj.get("description","")).lower() for kw in role_keywords)
    related_projects = [p for p in projects if project_related(p)]
    if len(projects)==0: report["Projects"]="❌ Missing"
    elif len(related_projects)==0: report["Projects"]="❌ Present but not relevant"
    elif len(projects)>=min_projects: report["Projects"], score="✅ Good", score+20
    else: report["Projects"], score="⚠️ Improve", score+(len(projects)/min_projects*20)

    # --- Contact Info ---
    contact_fields = ["name","email","phone","location"]
    missing = [f for f in contact_fields if not resume_data.get(f)]
    if not missing: report["Contact Info"], score="✅ Good", score+10
    else: report["Contact Info"], score="⚠️ Improve", score+(sum([1 for f in contact_fields if resume_data.get(f)])/len(contact_fields)*10)

    # --- Links ---
    links = resume_data.get("links", {})
    if links.get("linkedin") or links.get("github"): report["Links"], score="✅ Good", score+5
    else: report["Links"]="❌ Missing"

    # --- Certifications ---
    certifications = resume_data.get("certifications", [])
    cert_related = any(any(kw in cert.lower() for kw in role_keywords) for cert in certifications)
    if len(certifications)==0: report["Certifications"]="❌ Missing"
    elif not cert_related: report["Certifications"]="❌ Present but not relevant"
    else: report["Certifications"], score="✅ Good", score+5

    # --- Achievements ---
    achievements = resume_data.get("achievements", [])
    ach_related = any(any(kw in ach.lower() for kw in role_keywords) for ach in achievements)
    if len(achievements)==0: report["Achievements"]="❌ Missing"
    elif not ach_related: report["Achievements"]="❌ Present but not relevant"
    else: report["Achievements"], score="✅ Good", score+5

    # --- Role relevance ---
    if summary and any(k in summary.lower() for k in role_keywords): report["Role Relevance"], score="✅ Good", score+5
    else: report["Role Relevance"]="❌ Missing"

    return report, min(int(round(score)),100)

# ==============================
# 📌 Step 7: AI Feedback
# ==============================
def generate_ai_feedback(resume_data, role, candidate_type):
    prompt = f"""
You are an expert recruiter.
Evaluate this resume for role: {role}, Candidate type: {candidate_type.capitalize()}

Resume Data:
{json.dumps(resume_data)}

Guidelines:
1. Generate 9–12 concise bullet points, max 18 words each.
2. Focus on strengths, improvement areas, skills, projects, achievements, certifications.
3. Tailor suggestions strictly to role and candidate type.
4. Use action verbs.

Output:
- Bullet points only as plain text.
"""
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        st.warning(f"AI feedback generation failed: {e}")
        return "⚠️ AI feedback not available."

# ==============================
# 📌 Step 8: Streamlit UI
# ==============================
st.set_page_config(page_title="ATS Resume Checker", layout="wide")
st.title("📄 ATS Resume Checker")

resume_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
roles_dict = load_categories_roles_skills_tools_from_json("Technical_Roles.json")

# Category selection
categories = list(roles_dict.keys())
selected_category = st.selectbox("Select Category", categories) if categories else None
roles_under_category = list(roles_dict[selected_category].keys()) if selected_category else []
selected_role = st.selectbox("Select Role", roles_under_category) if roles_under_category else None

# Submit button
if st.button("✅ Submit"):
    if resume_file and selected_role:
        with st.spinner("Analyzing resume..."):
            resume_text, links_from_file = load_resume(resume_file)
            resume_data = parse_resume_with_gemini(resume_text, links_from_file)
            candidate_type = detect_candidate_type(resume_data, resume_text)
            sections_report, ats_score = evaluate_resume_sections_and_score(
                resume_data, roles_dict[selected_category][selected_role], candidate_type
            )
            ai_feedback = generate_ai_feedback(resume_data, selected_role, candidate_type)

        # Display results
        st.subheader(f"Candidate Type: {candidate_type.capitalize()}")
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("📌 ATS Score")
            st.write("**Total ATS Score:**", ats_score)
            st.progress(min(ats_score / 100, 1.0))
            st.subheader("📌 Section-wise ATS Check")
            for section, status in sections_report.items():
                color = "green" if "Good" in status else "orange" if "Improve" in status else "red"
                st.markdown(f"- **{section}**: <span style='color:{color}'>{status}</span>", unsafe_allow_html=True)

        with col2:
            st.subheader("📋 AI Feedback")
            st.markdown(ai_feedback.replace("-", "•"), unsafe_allow_html=True)

        # Download JSON report
        st.download_button(
            "📥 Download Report (JSON)",
            json.dumps({"sections": sections_report, "score": ats_score, "feedback": ai_feedback}, indent=2),
            "ats_report.json",
            "application/json"
        )
    else:
        st.warning("Please upload a resume and select a role before submitting.")
